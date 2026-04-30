# pip install streamlit sympy matplotlib pytesseract pillow python-docx numpy


import streamlit as st
from sympy import symbols, integrate, latex, lambdify, sympify
from sympy.parsing.sympy_parser import parse_expr, standard_transformations, implicit_multiplication_application
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
import io

# Configurazione Pagina
st.set_page_config(page_title="Risolutore Integrali Pro", layout="centered")

# --- TRASFORMAZIONI E CALCOLI ---
transformations = (standard_transformations + (implicit_multiplication_application,))

def pulisci_input(testo):
    """Converte il simbolo ^ in ** per il calcolo Python"""
    return testo.replace('^', '**')

def ottieni_suggerimento(f):
    """Fornisce suggerimenti teorici basati sulla funzione"""
    f_str = str(f).lower()
    if 'log' in f_str or 'exp' in f_str:
        return "Suggerimento: Per funzioni logaritmiche o esponenziali miste, prova l'integrazione per parti."
    if '/' in f_str:
        return "Suggerimento: Per funzioni fratte, verifica se il numeratore è la derivata del denominatore o usa i fratti semplici."
    if 'sin' in f_str or 'cos' in f_str:
        return "Suggerimento: Per funzioni trigonometriche, considera le identità fondamentali o la sostituzione."
    return "Suggerimento: Applica le regole di integrazione immediata (potenze, somma, costanti)."

# --- EXPORT FUNCTIONS ---
def crea_word(tipo, f_str, res_str, steps, fig):
    doc = Document()
    doc.add_heading(f'Esercizio Integrale {tipo}', 0)
    doc.add_heading('Testo:', level=1)
    doc.add_paragraph(f"Calcolare l'integrale della funzione: f(x) = {f_str}")
    doc.add_heading('Svolgimento:', level=1)
    for s in steps:
        doc.add_paragraph(s)
    doc.add_heading('Risultato Finale:', level=1)
    doc.add_paragraph(res_str)
    
    img_stream = io.BytesIO()
    fig.savefig(img_stream, format='png')
    doc.add_picture(img_stream, width=Inches(4))
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

def crea_pdf(tipo, f_str, res_str, steps, sug):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, f"Risoluzione Integrale {tipo}", ln=True, align='C')
    pdf.ln(10)
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(200, 10, f"Funzione: f(x) = {f_str}", ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", '', 11)
    for s in steps:
        pdf.multi_cell(0, 10, s)
    pdf.ln(5)
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(200, 10, f"Risultato: {res_str}", ln=True)
    pdf.ln(10)
    pdf.set_font("Arial", 'I', 10)
    pdf.set_text_color(50, 50, 200)
    pdf.multi_cell(0, 10, sug)
    return pdf.output()

# --- INTERFACCIA STREAMLIT ---
st.title("🧮 Integrals Master 2026")

if "input_math" not in st.session_state:
    st.session_state.input_math = ""

# Griglia Pulsanti
st.subheader("Inserimento Rapido")
c_btns = st.columns(8)
tasti = [("x", "x"), ("^", "^"), ("√", "sqrt("), ("/", "/"), ("sin", "sin("), ("cos", "cos("), ("ln", "log("), ("(", "(")]

for i, (lab, val) in enumerate(tasti):
    if c_btns[i].button(lab):
        st.session_state.input_math += val

# Input Testuale
testo_equazione = st.text_input("Equazione:", value=st.session_state.input_math)

# Tasto CANCELLA TUTTO
if st.button("🗑️ Cancella tutto", use_container_width=True):
    st.session_state.input_math = ""
    st.rerun()

# Limiti
col_a, col_b = st.columns(2)
a_in = col_a.text_input("Limite a")
b_in = col_b.text_input("Limite b")

# --- ELABORAZIONE ---
if st.button("RISOLVI E GENERA REPORT", type="primary", use_container_width=True):
    if testo_equazione:
        try:
            x = symbols('x')
            f_puro = pulisci_input(testo_equazione)
            f = parse_expr(f_puro, transformations=transformations)
            suggerimento = ottieni_suggerimento(f)
            
            # Calcolo
            if not a_in and not b_in:
                tipo = "Indefinito"
                res = integrate(f, x)
                res_txt = f"{res} + C"
                steps = [f"1. Analisi della funzione f(x) = {f}", "2. Applicazione delle regole di integrazione indefinita."]
                st.latex(rf"\int {latex(f)} \, dx = {latex(res)} + C")
            else:
                tipo = "Definito"
                a, b = sympify(pulisci_input(a_in)), sympify(pulisci_input(b_in))
                primitiva = integrate(f, x)
                res = integrate(f, (x, a, b))
                res_txt = str(res)
                steps = [
                    f"1. Funzione f(x) = {f}",
                    f"2. Calcolo della primitiva F(x) = {primitiva}",
                    f"3. Teorema fondamentale: F({b}) - F({a})"
                ]
                st.latex(rf"\int_{{{latex(a)}}}^{{{latex(b)}}} {latex(f)} \, dx = {latex(res)}")

            st.info(suggerimento)

            # Grafico
            fig, ax = plt.subplots()
            f_n = lambdify(x, f, "numpy")
            x_v = np.linspace(-10, 10, 400)
            ax.plot(x_v, f_n(x_v))
            ax.axhline(0, color='black', lw=1)
            ax.grid(True, alpha=0.3)
            st.pyplot(fig)

            # Download Buttons
            st.subheader("Esporta Risultati")
            down_col1, down_col2 = st.columns(2)
            
            word_data = crea_word(tipo, str(f), res_txt, steps, fig)
            down_col1.download_button("📄 Scarica Word", word_data, "integrale.docx")
            
            pdf_data = crea_pdf(tipo, str(f), res_txt, steps, suggerimento)
            down_col2.download_button("📕 Scarica PDF", pdf_data, "integrale.pdf")

        except Exception as e:
            st.error(f"Errore: {e}")


# streamlit run integrali.py
# per avviare il programma dal terminale nella stessa cartella
